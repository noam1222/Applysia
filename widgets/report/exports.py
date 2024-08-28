import os
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import matplotlib.animation as animation
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt

from constants import TIME_DB, APPLYSIA_DB, ALL_APPLYSIAS


def _get_time_from_format(time):
    return str(time).split(" ")[1][:-3]

def export_report_to_excel(report, file_path):

    # Initialize two lists to hold keys and values separately
    keys = []
    values = []

    # Extract keys and values from report dictionary
    for key, value in report.items():
        if key == "_id":
            continue
        keys.append(key)
        if key == TIME_DB:
            value = _get_time_from_format(value)
        elif key == APPLYSIA_DB and value == ALL_APPLYSIAS:
            value = "all"
        values.append(value)

    # Create DataFrame with keys and values as separate columns
    df = pd.DataFrame({'Key': keys, 'Value': values})

    df.to_excel(file_path, index=False)


def export_report_to_word(report, path):
    doc = Document()

    # Add a table with 2 columns
    table = doc.add_table(rows=0, cols=2)

    # Set table style and alignment
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Add table headers
    hdr_cells = table.add_row().cells
    hdr_cells[0].text = 'Key'
    hdr_cells[1].text = 'Value'

    # Populate the table with report data
    for key, value in report.items():
        if key == "_id":
            continue
        row_cells = table.add_row().cells
        row_cells[0].text = key
        if key == TIME_DB:
            value = _get_time_from_format(str(value))
        elif key == APPLYSIA_DB and value == ALL_APPLYSIAS:
            value = "all"
        row_cells[1].text = str(value)

    # Adjust table cell widths
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(10)

    # Save the document
    doc.save(path)


def draw_path_video(report, output_file, fps=10, smooth_factor=1000):
    """
    Draws the path made by an array of points and saves it as a video.

    :param points: List or array of points, where each point is a tuple (x, y).
    :param output_file: The name of the output video file.
    :param fps: Frames per second for the video.
    :param smooth_factor: Number of interpolation points between each given point to smooth the animation.
    """

    # Extract trail points and applysia number from the report
    points = report['trail_points']
    applysia_num = report['applysia']
    # Create a figure and axis
    fig, ax = plt.subplots()

    # Extract X and Y coordinates from the points
    x_data = [p['x'] for p in points]
    y_data = [p['y'] for p in points]

    # Interpolate the points to create a smoother line
    t = np.linspace(0, 1, len(points))
    t_new = np.linspace(0, 1, len(points) * smooth_factor)
    x_smooth = np.interp(t_new, t, x_data)
    y_smooth = np.interp(t_new, t, y_data)

    # Set up the plot limits
    ax.set_xlim(min(x_smooth) - 1, max(x_smooth) + 1)
    ax.set_ylim(min(y_smooth) - 1, max(y_smooth) + 1)

    # Remove numbers on x and y axes
    ax.set_xticks([])
    ax.set_yticks([])

    # Set the title of the plot
    ax.set_title(f"Path of Applysia {applysia_num}")

    # Initialize a line object
    line, = ax.plot([], [], lw=2)

    # Function to initialize the animation
    def init():
        line.set_data([], [])
        return line,

    # Function to update the line at each frame
    def update(frame):
        line.set_data(x_data[:frame], y_data[:frame])
        return line,

    # Create the animation
    ani = animation.FuncAnimation(fig, update, frames=len(points) + 1,
                                  init_func=init, blit=True)

    # Save the animation as a video
    ani.save(output_file, fps=fps)
    plt.show()

    plt.close()

